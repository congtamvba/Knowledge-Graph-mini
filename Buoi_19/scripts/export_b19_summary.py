from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = PROJECT_ROOT / "outputs" / "tong_ket_buoi_19.docx"


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_result_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Hạng mục", "Kết quả", "Minh chứng")
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = (
        ("Prompt Setup", "PASS", "Docker 29.7.2; Compose 5.4.0; dữ liệu và .env sẵn sàng."),
        ("Prompt 1 - Ollama Adapter", "PASS", "OllamaClient gọi /api/tags và /api/generate; có fallback an toàn."),
        ("Prompt 2 - Core Engines", "PASS", "UC1-UC4 hỗ trợ Ollama/Gemini; UC3 và UC4 đã chạy local."),
        ("Prompt 3 - Docker", "PASS", "Dockerfile, docker-compose.yml, requirements.txt và Streamlit app hợp lệ."),
        ("Prompt 4 - Runtime", "PASS", "Hai container Up; qwen3:0.6b 522 MB; Streamlit health 200/ok."),
        ("Prompt 5 - Security", "PASS", "6/6 kiểm tra đạt: RBAC, citation, review, audit privacy và local model."),
        ("Prompt 6 - Final Validation", "PASS", "LOCAL AI SYSTEM READY: YES."),
    )
    for result in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, result):
            cell.text = value


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    title = document.add_heading("TỔNG KẾT THỰC HÀNH BUỔI 19", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Đóng gói Local AI System với Docker, Ollama Qwen3:0.6B và Streamlit")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("1. Mục tiêu hoàn thành", level=1)
    document.add_paragraph(
        "Hệ thống RAG Bảo mật và Kiểm toán Agribank đã được chuyển sang Local AI với Ollama, "
        "đóng gói bằng Docker Compose và cung cấp dashboard Streamlit chạy tại localhost."
    )
    document.add_paragraph("Thành phần chính:", style="List Bullet")
    for item in (
        "Ollama Server chạy model qwen3:0.6b trên cổng 11434.",
        "Agribank Local AI Streamlit App chạy trên cổng 8501.",
        "Ollama API Adapter hỗ trợ health check, generate JSON và fallback an toàn.",
        "UC1 Internal Lookup, UC2 Compliance Gap, UC3 Compliance Checker và UC4 Audit Checklist.",
        "RBAC, citation integrity, NEEDS_HUMAN_REVIEW và audit log privacy.",
    ):
        document.add_paragraph(item, style="List Bullet 2")

    document.add_heading("2. Kết quả thực hiện", level=1)
    add_result_table(document)

    document.add_heading("3. Trạng thái nghiệm thu", level=1)
    add_code(
        document,
        "OLLAMA SERVER STATUS: PASS\n"
        "LOCAL MODEL QWEN3: PASS\n"
        "DOCKER CONTAINERIZATION: PASS\n"
        "LOCAL COMPLIANCE ENGINES: PASS\n\n"
        "LOCAL AI SYSTEM READY: YES",
    )

    document.add_heading("4. Kiến trúc vận hành", level=1)
    document.add_paragraph(
        "Browser truy cập Streamlit App tại http://localhost:8501. App dùng OllamaClient để gọi "
        "Ollama REST API tại http://ollama:11434 trong Docker network. Ollama dùng model qwen3:0.6b "
        "được lưu bền vững trong volume ollama_data."
    )

    document.add_heading("5. Kết quả kiểm thử bảo mật", level=1)
    for item in (
        "Prompt routing: LLM_PROVIDER=ollama đưa prompt đến Ollama local.",
        "RBAC: role Staff bị chặn 393 chunks hạn chế, không phát hiện dữ liệu rò rỉ.",
        "Citation Integrity: UC3 và UC4 chỉ xuất citation có trong corpus nguồn.",
        "Human Review: toàn bộ finding/checklist có review_status = NEEDS_HUMAN_REVIEW.",
        "Audit Log Privacy: log không chứa API key, password hoặc token; query nhạy cảm được redaction.",
        "Local Model Resilience: qwen3:0.6b phản hồi từ Docker volume local.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("6. Lệnh demo", level=1)
    add_code(document, "docker compose ps")
    add_code(document, "docker exec agribank-ollama-server ollama list")
    add_code(document, "python scripts\\compliance_checker.py")
    add_code(document, "python scripts\\audit_checklist_gen.py")
    add_code(document, "python scripts\\verify_b19_docker.py")
    document.add_paragraph("Mở UI: http://localhost:8501. Chọn vai trò Kiểm toán viên, sau đó chạy UC3 và UC4.")

    document.add_heading("7. Artifact nộp bài", level=1)
    for item in (
        "outputs/b19_docker_acceptance_report.md",
        "outputs/b19_security_test_report.md",
        "outputs/compliance_conflicts.csv",
        "outputs/audit_checklist_results.csv",
        "outputs/audit_log.jsonl",
        "outputs/tong_ket_buoi_19.docx",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_paragraph("Ghi chú: Air-gapped vật lý cần được xác nhận bằng thao tác tắt Wi-Fi/rút mạng tại máy host trong buổi demo.")
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(f"WORD_SUMMARY_CREATED={OUTPUT_PATH}")
